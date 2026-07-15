"""Engineering Copilot — Modulo 4 (Document Generator). Gera os 8 entregaveis
do MVP em DOCX a partir dos dados ja extraidos (Modulos 1/2) e da avaliacao de
compliance (Modulo 5). Mesmo padrao do gerar_pgr_docx do NR1: marca d'agua de
texto quando a licenca nao e premium."""
import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

WATERMARK_TEXT = "DEMONSTRACAO — VERSAO NAO LICENCIADA"
CHECKOUT_URL = "https://buy.stripe.com/5kQ7sN4ModwDblO6q4g7e0l"

TITULOS = {
    "memorial": "Memorial Descritivo",
    "relatorio_tecnico": "Relatorio Tecnico",
    "inventario": "Inventario de Equipamentos",
    "checklist": "Checklist Tecnico de Conformidade",
    "plano_manutencao": "Plano de Manutencao",
    "relatorio_fotografico": "Relatorio Fotografico",
    "data_book": "Data Book",
    "relatorio_executivo": "Relatorio Executivo",
    "as_built": "Documentacao As Built",
}


def _cabecalho(document: Document, titulo: str, projeto: dict, licenca_premium: bool) -> None:
    h = document.add_heading(titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Cliente: {projeto.get('cliente', '')}")
    document.add_paragraph(f"Local: {projeto.get('local', '')}")
    document.add_paragraph(f"Data de emissao: {date.today().strftime('%d/%m/%Y')}")
    if not licenca_premium:
        p = document.add_paragraph()
        run = p.add_run(WATERMARK_TEXT + " — documento nao valido para fins legais ate licenciamento.")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x1E, 0x1E)
        p2 = document.add_paragraph()
        run2 = p2.add_run(f"Assine para remover a marca d'agua e liberar todos os documentos: {CHECKOUT_URL}")
        run2.italic = True
        run2.font.color.rgb = RGBColor(0x0D, 0x2C, 0x4C)
    document.add_paragraph()


def _secao_equipamentos(document: Document, equipamentos: list[dict]) -> None:
    document.add_heading("Equipamentos", level=1)
    if not equipamentos:
        document.add_paragraph("Nenhum equipamento identificado ate o momento.")
        return
    tabela = document.add_table(rows=1, cols=5)
    tabela.style = "Light Grid Accent 1"
    for cel, texto in zip(tabela.rows[0].cells, ["TAG", "Tipo", "Fabricante", "Modelo", "Localizacao"]):
        cel.text = texto
    for eq in equipamentos:
        row = tabela.add_row().cells
        row[0].text = eq.get("tag") or "-"
        row[1].text = eq.get("tipo") or "-"
        row[2].text = eq.get("fabricante") or "-"
        row[3].text = eq.get("modelo") or "-"
        row[4].text = eq.get("localizacao") or "-"


def _secao_normas(document: Document, normas: list[dict]) -> None:
    document.add_heading("Normas Aplicaveis", level=1)
    if not normas:
        document.add_paragraph("Nenhuma norma identificada ate o momento.")
        return
    for n in normas:
        document.add_paragraph(f"• {n.get('norma', '')} — {n.get('descricao', '')}", style="List Bullet")


def _secao_compliance(document: Document, compliance: dict) -> None:
    document.add_heading("Conformidade", level=1)
    document.add_paragraph(f"Score de conformidade: {compliance.get('score_conformidade', 0)}/100")
    if compliance.get("conformidades"):
        document.add_paragraph("Itens conformes:")
        for c in compliance["conformidades"]:
            document.add_paragraph(f"• {c['norma']} — {c['descricao']}", style="List Bullet")
    if compliance.get("pendencias"):
        document.add_paragraph("Pendencias:")
        for p in compliance["pendencias"]:
            document.add_paragraph(f"• {p['norma']} — {p['descricao']}", style="List Bullet")


def _secao_fotos(document: Document, fotos: list[dict]) -> None:
    document.add_heading("Registro Fotografico", level=1)
    if not fotos:
        document.add_paragraph("Nenhuma foto processada ate o momento.")
        return
    for f in fotos:
        analise = f.get("analise") or {}
        document.add_paragraph(f"Foto: {f.get('nome_arquivo', '')}", style="Intense Quote")
        if analise.get("tag_ocr"):
            document.add_paragraph(f"TAG identificada (OCR): {analise['tag_ocr']}")
        if analise.get("problemas"):
            document.add_paragraph("Problemas detectados: " + ", ".join(analise["problemas"]))
        if analise.get("recomendacoes"):
            document.add_paragraph("Recomendacoes: " + ", ".join(analise["recomendacoes"]))


def _secao_documentos(document: Document, documentos: list[dict]) -> None:
    document.add_heading("Documentos Recebidos", level=1)
    if not documentos:
        document.add_paragraph("Nenhum documento recebido ate o momento.")
        return
    for d in documentos:
        document.add_paragraph(f"• {d.get('nome_arquivo', '')}", style="List Bullet")


def _plano_manutencao_por_norma(compliance: dict) -> list[str]:
    periodicidade = {
        "ABNT NBR 13971": "Verificacao mensal (limpeza de filtros/bandejas) e relatorio PMOC trimestral.",
        "NR-13": "Inspecao de seguranca periodica conforme categoria do vaso/caldeira (1 a 5 anos).",
        "NR-12": "Inspecao e manutencao preventiva conforme manual do fabricante e NR-12.",
        "NR-10": "Inspecao eletrica anual e verificacao de aterramento.",
        "NR-35": "Verificacao de ancoragem e EPIs antes de cada trabalho em altura.",
        "ART/CREA": "Renovar ART a cada nova intervencao tecnica relevante.",
    }
    linhas = []
    for item in compliance.get("conformidades", []) + compliance.get("pendencias", []):
        norma = item.get("norma", "")
        if norma in periodicidade:
            linhas.append(f"{norma}: {periodicidade[norma]}")
    return linhas or ["Nenhuma recomendacao de periodicidade aplicavel com os dados atuais."]


def _corpo_memorial(document, ctx):
    document.add_heading("Escopo", level=1)
    document.add_paragraph(ctx["projeto"].get("escopo") or "Escopo nao informado.")
    _secao_equipamentos(document, ctx["equipamentos"])
    _secao_normas(document, ctx["normas"])


def _corpo_relatorio_tecnico(document, ctx):
    _secao_documentos(document, ctx["documentos"])
    _secao_equipamentos(document, ctx["equipamentos"])
    _secao_compliance(document, ctx["compliance"])


def _corpo_inventario(document, ctx):
    _secao_equipamentos(document, ctx["equipamentos"])


def _corpo_checklist(document, ctx):
    document.add_heading("Checklist de Conformidade", level=1)
    for item in ctx["compliance"].get("conformidades", []) + ctx["compliance"].get("pendencias", []):
        marca = "[OK]" if item["status"] == "conforme" else "[PENDENTE]"
        document.add_paragraph(f"{marca} {item['norma']} — {item['descricao']}", style="List Bullet")


def _corpo_plano_manutencao(document, ctx):
    document.add_heading("Plano de Manutencao", level=1)
    for linha in _plano_manutencao_por_norma(ctx["compliance"]):
        document.add_paragraph(f"• {linha}", style="List Bullet")


def _corpo_relatorio_fotografico(document, ctx):
    _secao_fotos(document, ctx["fotos"])


def _corpo_data_book(document, ctx):
    _secao_documentos(document, ctx["documentos"])
    _secao_equipamentos(document, ctx["equipamentos"])
    _secao_normas(document, ctx["normas"])
    _secao_fotos(document, ctx["fotos"])
    _secao_compliance(document, ctx["compliance"])


def _corpo_relatorio_executivo(document, ctx):
    compliance = ctx["compliance"]
    document.add_heading("Resumo Executivo", level=1)
    document.add_paragraph(
        f"Projeto: {ctx['projeto'].get('cliente', '')} — {len(ctx['equipamentos'])} equipamento(s) "
        f"identificado(s), {len(ctx['normas'])} norma(s) mapeada(s), score de conformidade "
        f"de {compliance.get('score_conformidade', 0)}/100."
    )
    if compliance.get("pendencias"):
        document.add_paragraph("Principais pendencias:")
        for p in compliance["pendencias"][:5]:
            document.add_paragraph(f"• {p['norma']} — {p['descricao']}", style="List Bullet")


def _corpo_as_built(document, ctx):
    document.add_heading("Documentacao As Built", level=1)
    document.add_paragraph(
        "Registro do estado final instalado do projeto, consolidado a partir dos "
        "documentos recebidos, fotos de campo e verificacao de conformidade — "
        "referencia para operacao, manutencao e futuras intervencoes."
    )
    _secao_equipamentos(document, ctx["equipamentos"])
    _secao_fotos(document, ctx["fotos"])
    document.add_heading("Desvios Identificados", level=1)
    pendencias = ctx["compliance"].get("pendencias", [])
    if not pendencias:
        document.add_paragraph("Nenhum desvio de conformidade identificado ate o momento.")
    else:
        for p in pendencias:
            document.add_paragraph(f"• {p['norma']} — {p['descricao']}", style="List Bullet")


_CONSTRUTORES = {
    "memorial": _corpo_memorial,
    "relatorio_tecnico": _corpo_relatorio_tecnico,
    "inventario": _corpo_inventario,
    "checklist": _corpo_checklist,
    "plano_manutencao": _corpo_plano_manutencao,
    "relatorio_fotografico": _corpo_relatorio_fotografico,
    "data_book": _corpo_data_book,
    "relatorio_executivo": _corpo_relatorio_executivo,
    "as_built": _corpo_as_built,
}


def gerar_documento(tipo: str, projeto: dict, equipamentos: list[dict], normas: list[dict],
                     documentos: list[dict], fotos: list[dict], compliance: dict,
                     licenca_premium: bool = False) -> io.BytesIO:
    if tipo not in _CONSTRUTORES:
        raise ValueError(f"Tipo de documento desconhecido: {tipo}. Opcoes: {sorted(_CONSTRUTORES)}")

    document = Document()
    _cabecalho(document, TITULOS[tipo], projeto, licenca_premium)
    _CONSTRUTORES[tipo](document, {
        "projeto": projeto, "equipamentos": equipamentos, "normas": normas,
        "documentos": documentos, "fotos": fotos, "compliance": compliance,
    })

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf
